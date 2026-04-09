"""
╔══════════════════════════════════════════════════════════════════════════════╗
║  DKB POWER ENGINEERING CO., LTD.                                            ║
║  Company Profile Dashboard  v3.0  —  dkb_profile_dashboard.py               ║
║                                                                              ║
║  Fixes in v3:                                                               ║
║  • Thai font rendering: w:rFonts with ascii+hAnsi+cs+eastAsia on every run  ║
║  • No placeholders anywhere in the generated .docx                          ║
║  • use_container_width → width (Streamlit 1.40+ compatible)                 ║
║  • Dynamic custom sections (Preventive Maintenance, Construction, etc.)     ║
║  • Spacious, elegant docx layout with proper white space                    ║
║                                                                              ║
║  Run:  streamlit run dkb_profile_dashboard.py                               ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""

import io, json, os, re, uuid
from copy import deepcopy
from datetime import datetime
from pathlib import Path

import streamlit as st
from PIL import Image, ImageOps

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ══════════════════════════════════════════════════════════════════════════════
#  PATHS
# ══════════════════════════════════════════════════════════════════════════════
BASE_DIR     = Path(__file__).parent
PHOTOS_DIR   = BASE_DIR / "photos"
OUTPUT_DIR   = BASE_DIR / "output"
ASSETS_DIR   = BASE_DIR / "assets"
DATA_FILE    = BASE_DIR / "projects.json"
COMPANY_FILE = BASE_DIR / "company.json"
SECTIONS_FILE= BASE_DIR / "custom_sections.json"

for _d in (PHOTOS_DIR, OUTPUT_DIR, ASSETS_DIR):
    _d.mkdir(exist_ok=True)


# ══════════════════════════════════════════════════════════════════════════════
#  DEFAULTS
# ══════════════════════════════════════════════════════════════════════════════
DEFAULT_COMPANY = {
    "name_th"     : "บริษัท ดีเคบี เพาเวอร์ เอนจิเนียริ่ง จำกัด",
    "name_en"     : "DKB POWER ENGINEERING CO., LTD.",
    "reg_no"      : "0105565009633",
    "reg_date_th" : "17 มกราคม 2565",
    "reg_date_en" : "17 January 2022",
    "capital_th"  : "1,000,000.00 บาท",
    "capital_en"  : "THB 1,000,000.00",
    "address_th"  : "422/635/1 ซอยสุจินทวงศ์ 34 แขวงเสน่แสบ เขตมีนบุรี กรุงเทพมหานคร",
    "address_en"  : "422/635/1 Soi Sujinwong 34, Saen Saeb, Min Buri, Bangkok",
    "tel"         : "082-234-4680",
    "email"       : "kasama.D.K.B@gmail.com",
    "year"        : "2026",
    "history_th"  : (
        "บริษัท ดีเคบี เพาเวอร์ เอนจิเนียริ่ง จำกัด ก่อตั้งขึ้นด้วยเจตนารมณ์อันแน่วแน่ "
        "ในการให้บริการด้านวิศวกรรมระบบไฟฟ้าและสื่อสารอย่างครบวงจร มุ่งเน้นงานติดตั้ง"
        "ระบบไฟฟ้า ระบบเสาอากาศโทรทัศน์ ระบบกล้องวงจรปิด (CCTV) ระบบทีวีดาวเทียม "
        "และระบบทีวีดิจิทัล ให้แก่ลูกค้าทั้งภาครัฐและภาคเอกชน ด้วยทีมช่างผู้ชำนาญการ "
        "ที่มุ่งมั่นส่งมอบงานคุณภาพสูง ตรงเวลา และอยู่ภายในงบประมาณที่ตกลงไว้"
    ),
    "history_en"  : (
        "DKB Power Engineering Co., Ltd. was established with an unwavering commitment "
        "to delivering comprehensive electrical and communication engineering services. "
        "Specialising in electrical installations, TV antenna systems, CCTV, satellite TV, "
        "and digital TV, the company serves both public and private sector clients. "
        "Our certified technical team upholds the highest engineering standards, "
        "delivering every project on time and within budget."
    ),
    "logo_file"   : "",
}

DEFAULT_PROJECT = {
    "id"          : "proj_001",
    "no"          : 1,
    "name_th"     : "โครงการเรือนแถวตำรวจ สภ.บางโทรัด จ.สมุทรสาคร",
    "name_en"     : "Police Row House Project, Bang Thorat Police Station, Samut Sakhon",
    "customer_th" : "สภ.บางโทรัด จ.สมุทรสาคร",
    "customer_en" : "Bang Thorat Police Station, Samut Sakhon Province",
    "main_contact": "บ.ณัฐกิจการช่าง",
    "location"    : "จ.สมุทรสาคร / Samut Sakhon",
    "desc_th"     : "งานติดตั้งระบบทีวี",
    "desc_en"     : "TV System Installation",
    "period"      : "2567",
    "value"       : "35,310.00",
    "photos"      : [],
    "created_at"  : "2026-01-01T00:00:00",
}

# Colour hex strings (used both in XML helpers and RGBColor)
C = {
    "navy"     : "0A3C7A",
    "blue"     : "2979C8",
    "darknavy" : "062A5A",
    "ltblue"   : "EDF4FC",
    "rowalt"   : "F4F8FC",
    "white"    : "FFFFFF",
    "midgray"  : "607A99",
    "textdark" : "1A2E45",
    "steelblue": "90B8E0",
    "paleblue" : "D0E4F5",
}

def _rgb(key: str) -> RGBColor:
    h = C[key]
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))


# ══════════════════════════════════════════════════════════════════════════════
#  DATA LAYER
# ══════════════════════════════════════════════════════════════════════════════

def load_company() -> dict:
    if COMPANY_FILE.exists():
        try:
            return {**DEFAULT_COMPANY, **json.loads(COMPANY_FILE.read_text("utf-8"))}
        except Exception:
            pass
    return deepcopy(DEFAULT_COMPANY)

def save_company(d: dict):
    COMPANY_FILE.write_text(json.dumps(d, ensure_ascii=False, indent=2), "utf-8")

def load_projects() -> list:
    if not DATA_FILE.exists():
        save_projects([deepcopy(DEFAULT_PROJECT)])
        return [deepcopy(DEFAULT_PROJECT)]
    try:
        data = json.loads(DATA_FILE.read_text("utf-8"))
        return data if isinstance(data, list) else []
    except Exception:
        return []

def save_projects(p: list):
    DATA_FILE.write_text(json.dumps(p, ensure_ascii=False, indent=2), "utf-8")

def load_sections() -> list:
    """Custom sections (Preventive Maintenance, Construction, etc.)"""
    if SECTIONS_FILE.exists():
        try:
            data = json.loads(SECTIONS_FILE.read_text("utf-8"))
            return data if isinstance(data, list) else []
        except Exception:
            pass
    return []

def save_sections(s: list):
    SECTIONS_FILE.write_text(json.dumps(s, ensure_ascii=False, indent=2), "utf-8")

def next_no(projects: list) -> int:
    return max((p.get("no", 0) for p in projects), default=0) + 1

def safe_fname(original: str) -> str:
    stem = re.sub(r"[^\w\-.]", "_", Path(original).stem)[:40]
    return f"{uuid.uuid4().hex[:8]}_{stem}{Path(original).suffix.lower()}"

def save_photos(files) -> list:
    saved = []
    for f in (files or []):
        name = safe_fname(f.name)
        (PHOTOS_DIR / name).write_bytes(f.read())
        saved.append(name)
    return saved

def save_logo(file) -> str:
    name = safe_fname(file.name)
    (ASSETS_DIR / name).write_bytes(file.read())
    return name


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX — FONT HELPER  (THE KEY FIX FOR THAI)
# ══════════════════════════════════════════════════════════════════════════════

THAI_FONT = "TH Sarabun New"
LATIN_FONT = "Calibri"

def _apply_font(run, font_name: str, size_pt: float,
                bold=False, italic=False, color_key: str = None):
    """
    Apply font to a run with ALL four w:rFonts attributes set.
    This is the critical fix — python-docx's run.font.name only sets
    w:ascii/hAnsi but NOT w:cs (complex script) which Thai requires.
    Setting all four prevents the '?' substitution in Word.
    """
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color_key:
        run.font.color.rgb = _rgb(color_key)

    # Build / replace w:rFonts with all 4 attributes
    rPr = run._r.get_or_add_rPr()
    # Remove any existing rFonts
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),    font_name)
    rFonts.set(qn("w:hAnsi"),    font_name)
    rFonts.set(qn("w:cs"),       font_name)   # ← Thai / Arabic / etc.
    rFonts.set(qn("w:eastAsia"), font_name)
    rPr.insert(0, rFonts)


def _run_th(para, text: str, size=10, bold=False, italic=False, color_key=None):
    """Add a Thai run."""
    r = para.add_run(text)
    _apply_font(r, THAI_FONT, size, bold, italic, color_key)
    return r

def _run_en(para, text: str, size=9, bold=False, italic=False, color_key=None):
    """Add a Latin/English run."""
    r = para.add_run(text)
    _apply_font(r, LATIN_FONT, size, bold, italic, color_key)
    return r


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX — XML HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _cell_bg(cell, hex6: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6)
    tcPr.append(shd)

def _cell_borders(cell, color="D8E6F2", size=4):
    tcPr    = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"),  str(size))
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)

def _cell_no_border(cell):
    tcPr    = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"),  "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tcPr.append(borders)

def _cell_margins(cell, top=80, bottom=80, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    mar  = OxmlElement("w:tcMar")
    for s, v in [("top",top),("bottom",bottom),("left",left),("right",right)]:
        el = OxmlElement(f"w:{s}")
        el.set(qn("w:w"), str(v))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)

def _col_width(cell, twips: int):
    tcPr = cell._tc.get_or_add_tcPr()
    w    = OxmlElement("w:tcW")
    w.set(qn("w:w"),    str(int(twips)))
    w.set(qn("w:type"), "dxa")
    tcPr.append(w)

def _para_space(p, before=0, after=0):
    pPr = p._p.get_or_add_pPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), str(before))
    spc.set(qn("w:after"),  str(after))
    pPr.append(spc)

def _bottom_rule(p, color="2979C8", size=8):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:color"), color)
    bot.set(qn("w:space"), "4")
    pBdr.append(bot)
    pPr.append(pBdr)

def _left_bar(p, color="2979C8", size=20):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    lft  = OxmlElement("w:left")
    lft.set(qn("w:val"),   "single")
    lft.set(qn("w:sz"),    str(size))
    lft.set(qn("w:color"), color)
    lft.set(qn("w:space"), "6")
    pBdr.append(lft)
    pPr.append(pBdr)

def _cp(cell, idx=0):
    while len(cell.paragraphs) <= idx:
        cell.add_paragraph()
    return cell.paragraphs[idx]

def _spacer(doc, pts=8):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:line"),     str(int(pts * 20)))
    spc.set(qn("w:lineRule"), "exact")
    spc.set(qn("w:before"),   "0")
    spc.set(qn("w:after"),    "0")
    pPr.append(spc)


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX — REUSABLE LAYOUT ATOMS
# ══════════════════════════════════════════════════════════════════════════════

def _section_header(doc, th: str, en: str):
    """Blue left-bar section header with Thai + English."""
    p = doc.add_paragraph()
    _para_space(p, before=160, after=60)
    _left_bar(p, color=C["navy"], size=22)
    _run_th(p, f"  {th}", size=11, bold=True, color_key="navy")
    _run_en(p, f"  /  {en.upper()}", size=8, color_key="midgray")
    _bottom_rule(p, color=C["blue"], size=6)


def _page_top_rule(doc):
    """4px navy top accent rule."""
    p = doc.add_paragraph()
    _para_space(p, before=0, after=0)
    _bottom_rule(p, color=C["navy"], size=28)


def _running_header(doc, co: dict):
    """Slim running company name below the top rule."""
    p = doc.add_paragraph()
    _para_space(p, before=40, after=80)
    _run_th(p, co["name_th"] + "  ", size=9, bold=True, color_key="navy")
    _run_en(p, co["name_en"], size=8, color_key="midgray")
    _bottom_rule(p, color=C["blue"], size=4)


def _footer_bar(doc, co: dict, page_num: int):
    """Dark navy contact footer."""
    _spacer(doc, 10)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    items  = [f"โทร  {co['tel']}", f"อีเมล  {co['email']}", f"หน้า / Page  {page_num}"]
    widths = [int(Cm(6.0).twips), int(Cm(8.0).twips), int(Cm(4.0).twips)]
    aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT]
    for j, (txt, w, aln) in enumerate(zip(items, widths, aligns)):
        c = tbl.rows[0].cells[j]
        _cell_bg(c, C["darknavy"])
        _cell_borders(c, color=C["navy"], size=4)
        _cell_margins(c, top=100, bottom=100, left=120, right=120)
        _col_width(c, w)
        p = _cp(c)
        p.alignment = aln
        _run_th(p, txt, size=8, color_key="white")


def _resize_image(path: Path, max_px=900) -> io.BytesIO:
    """Open, EXIF-rotate, downsize, return JPEG buffer."""
    img = Image.open(str(path)).convert("RGB")
    img = ImageOps.exif_transpose(img)
    if max(img.size) > max_px:
        img.thumbnail((max_px, max_px), Image.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85, optimize=True)
    buf.seek(0)
    return buf


def _photo_grid(doc, photo_paths: list, img_width_cm=4.0, cols=4):
    """Insert photos in a borderless grid table."""
    chunks = [photo_paths[i:i+cols] for i in range(0, len(photo_paths), cols)]
    for chunk in chunks:
        tbl = doc.add_table(rows=1, cols=cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        for j in range(cols):
            cell = tbl.rows[0].cells[j]
            _cell_no_border(cell)
            _cell_margins(cell, top=40, bottom=40, left=40, right=40)
            _col_width(cell, int(Cm(img_width_cm + 0.2).twips))
            pp = _cp(cell)
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if j < len(chunk):
                try:
                    buf = _resize_image(chunk[j])
                    pp.add_run().add_picture(buf, width=Cm(img_width_cm))
                except Exception:
                    _run_en(pp, f"[Photo {j+1}]", size=7, color_key="midgray")
            else:
                _cell_bg(cell, C["rowalt"])
        _spacer(doc, 6)


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX — PAGE BUILDERS
# ══════════════════════════════════════════════════════════════════════════════

def _build_cover(doc: Document, co: dict):
    _page_top_rule(doc)
    _spacer(doc, 18)

    # Logo
    logo_path = ASSETS_DIR / co.get("logo_file","") if co.get("logo_file") else None
    if logo_path and logo_path.exists():
        p = doc.add_paragraph()
        _para_space(p, before=0, after=80)
        p.add_run().add_picture(str(logo_path), height=Cm(1.6))

    # Company name
    p = doc.add_paragraph()
    _para_space(p, before=0, after=20)
    _run_th(p, co["name_th"], size=14, bold=True, color_key="navy")

    p = doc.add_paragraph()
    _para_space(p, before=0, after=60)
    _run_en(p, co["name_en"], size=10, bold=True, color_key="blue")

    # Big title
    p = doc.add_paragraph()
    _para_space(p, before=200, after=40)
    _run_en(p, "Company Profile", size=42, bold=True, color_key="navy")

    p = doc.add_paragraph()
    _para_space(p, before=0, after=30)
    _run_th(p, "โปรไฟล์บริษัท", size=16, color_key="midgray")

    _bottom_rule(doc.add_paragraph(), color=C["blue"], size=12)
    _spacer(doc, 12)

    # Descriptor
    p = doc.add_paragraph()
    _para_space(p, before=0, after=40)
    _run_th(p,
        "ผู้เชี่ยวชาญด้านวิศวกรรมระบบไฟฟ้า ระบบสื่อสาร ระบบโทรทัศน์ดิจิทัล "
        "ดาวเทียม และกล้องวงจรปิด\n",
        size=11, color_key="textdark")
    _run_en(p,
        "Specialist in Electrical Systems, TV Antenna, Digital TV, "
        "Satellite TV, CCTV & Communication Engineering",
        size=9, italic=True, color_key="midgray")

    # Service tags
    _spacer(doc, 10)
    p = doc.add_paragraph()
    _para_space(p, before=0, after=0)
    tags = ["Electrical Systems","TV Antenna","Digital TV","Satellite TV","CCTV","Communication Eng."]
    _run_en(p, "  ·  ".join(tags), size=9, color_key="blue")

    # Year watermark
    _spacer(doc, 80)
    p = doc.add_paragraph()
    _para_space(p, before=0, after=0)
    r = p.add_run(co["year"])
    _apply_font(r, LATIN_FONT, 52, bold=True,
                color_key=None)
    r.font.color.rgb = RGBColor(0xD8, 0xE8, 0xF5)

    _footer_bar(doc, co, 1)
    doc.add_page_break()


def _build_company_info(doc: Document, co: dict):
    _page_top_rule(doc)
    _running_header(doc, co)

    # Page title
    p = doc.add_paragraph()
    _para_space(p, before=20, after=60)
    _run_th(p, "ข้อมูลบริษัท  ", size=15, bold=True, color_key="navy")
    _run_en(p, "COMPANY INFORMATION", size=9, color_key="midgray")
    _bottom_rule(p, color=C["blue"], size=12)

    # Company name card
    nb = doc.add_table(rows=1, cols=1)
    c  = nb.rows[0].cells[0]
    _cell_bg(c, C["ltblue"])
    _cell_borders(c, color=C["blue"], size=8)
    _cell_margins(c, top=120, bottom=120, left=200, right=200)
    _col_width(c, int(Cm(18.0).twips))
    lp = _cp(c)
    _run_th(lp, co["name_th"] + "\n", size=14, bold=True, color_key="navy")
    _run_en(lp, co["name_en"],        size=10, bold=True, color_key="blue")

    _spacer(doc, 14)

    # Two-column body: registration (left) | history + services (right)
    body = doc.add_table(rows=1, cols=2)
    body.alignment = WD_TABLE_ALIGNMENT.LEFT
    lc = body.rows[0].cells[0]
    rc = body.rows[0].cells[1]
    for cx in (lc, rc): _cell_no_border(cx)
    _cell_margins(lc, 0, 0, 0, 120); _cell_margins(rc, 0, 0, 120, 0)
    _col_width(lc, int(Cm(8.5).twips)); _col_width(rc, int(Cm(9.0).twips))

    # ── Left: registration ──────────────────────────────────────────────────
    rh = lc.add_paragraph()
    _para_space(rh, before=0, after=40)
    _run_th(rh, "ข้อมูลการจดทะเบียน  ", size=9, bold=True, color_key="navy")
    _run_en(rh, "REGISTRATION DETAILS", size=7.5, color_key="midgray")
    _bottom_rule(rh, color=C["blue"], size=6)

    fields = [
        ("เลขทะเบียนนิติบุคคล\nRegistration No.",    co["reg_no"]),
        ("วันที่จดทะเบียน\nDate of Incorporation",    co["reg_date_th"] + "\n" + co["reg_date_en"]),
        ("ทุนจดทะเบียน\nRegistered Capital",          co["capital_th"]  + "\n" + co["capital_en"]),
        ("ที่ตั้งสำนักงาน\nOffice Address",
         co["address_th"] + "\n" + co["address_en"]),
        ("โทรศัพท์ / Tel", co["tel"]),
        ("อีเมล / Email",  co["email"]),
    ]
    reg = lc.add_table(rows=len(fields), cols=2)
    reg.style = "Table Grid"
    for i, (lbl, val) in enumerate(fields):
        lcc = reg.rows[i].cells[0]
        vcc = reg.rows[i].cells[1]
        _cell_bg(lcc, C["ltblue"] if i % 2 == 0 else "F4F8FC")
        _cell_bg(vcc, C["white"])
        _cell_borders(lcc, color="D8E6F2"); _cell_borders(vcc, color="D8E6F2")
        _cell_margins(lcc, 60,60,80,60); _cell_margins(vcc, 60,60,80,60)
        _col_width(lcc, int(Cm(3.2).twips)); _col_width(vcc, int(Cm(5.3).twips))
        lp2 = _cp(lcc)
        for k, ln in enumerate(lbl.split("\n")):
            _run_th(lp2, ln + ("\n" if k==0 and "\n" in lbl else ""),
                    size=8.5 if k==0 else 7,
                    bold=(k==0),
                    color_key="navy" if k==0 else "midgray")
        vp2 = _cp(vcc)
        for k, ln in enumerate(val.split("\n")):
            _run_th(vp2, ln + ("\n" if k < len(val.split("\n"))-1 else ""),
                    size=9 if k==0 else 7.5,
                    bold=(k==0),
                    color_key="textdark" if k==0 else "midgray")

    # ── Right: history ──────────────────────────────────────────────────────
    hh = rc.add_paragraph()
    _para_space(hh, before=0, after=40)
    _run_th(hh, "ประวัติและพันธกิจ  ", size=9, bold=True, color_key="navy")
    _run_en(hh, "HISTORY & MISSION", size=7.5, color_key="midgray")
    _bottom_rule(hh, color=C["blue"], size=6)

    ht = rc.add_table(rows=1, cols=1)
    hc = ht.rows[0].cells[0]
    _cell_bg(hc, "F4F8FC"); _cell_borders(hc, color=C["blue"], size=6)
    _cell_margins(hc, 100,100,140,140); _col_width(hc, int(Cm(9.0).twips))
    hp = _cp(hc)
    _run_th(hp, co["history_th"] + "\n\n", size=9, color_key="textdark")
    _run_en(hp, co["history_en"], size=8, italic=True, color_key="midgray")

    # ── Right: services grid ────────────────────────────────────────────────
    sp = rc.add_paragraph(); _para_space(sp, before=80, after=40)
    sh = rc.add_paragraph(); _para_space(sh, before=0, after=40)
    _run_th(sh, "ขอบเขตงานบริการ  ", size=9, bold=True, color_key="navy")
    _run_en(sh, "SCOPE OF SERVICES", size=7.5, color_key="midgray")
    _bottom_rule(sh, color=C["blue"], size=6)

    services = [
        ("ระบบไฟฟ้า",          "Electrical Installation"),
        ("ระบบกล้อง CCTV",     "CCTV System"),
        ("ระบบทีวีดิจิทัล",   "Digital TV System"),
        ("ทีวีดาวเทียม",       "Satellite TV"),
        ("เสาอากาศทีวี",       "TV Antenna System"),
        ("งานวิศวกรรมอื่นๆ",  "Other Engineering Works"),
    ]
    sg = rc.add_table(rows=3, cols=2); sg.alignment = WD_TABLE_ALIGNMENT.LEFT
    idx = 0
    for ri in range(3):
        for ci in range(2):
            if idx >= len(services): break
            sc = sg.rows[ri].cells[ci]
            _cell_bg(sc, C["ltblue"]); _cell_borders(sc, color="D8E6F2")
            _cell_margins(sc, 70,70,90,70); _col_width(sc, int(Cm(4.5).twips))
            scp = _cp(sc)
            _run_th(scp, services[idx][0]+"\n", size=8.5, bold=True, color_key="navy")
            _run_en(scp, services[idx][1], size=7.5, color_key="midgray")
            idx += 1

    _footer_bar(doc, co, 2)
    doc.add_page_break()


def _build_project_reference(doc: Document, co: dict, projects: list):
    _page_top_rule(doc)
    _running_header(doc, co)

    p = doc.add_paragraph()
    _para_space(p, before=20, after=60)
    _run_th(p, "รายการอ้างอิงโครงการ  ", size=15, bold=True, color_key="navy")
    _run_en(p, "PROJECT REFERENCE LIST", size=9, color_key="midgray")
    _bottom_rule(p, color=C["blue"], size=12)
    _spacer(doc, 8)

    cols_def = [
        ("ที่",               "No.",             0.85),
        ("ชื่อโครงการ",        "Project Name",    4.2),
        ("ผู้ว่าจ้าง / ลูกค้า","Customer / Owner", 3.4),
        ("ผู้รับช่วง",        "Main Contact",    2.7),
        ("สถานที่",           "Location",        2.5),
        ("รายละเอียดงาน",     "Job Description", 3.0),
        ("มูลค่า (บาท)",      "Value (THB)",     1.9),
    ]

    tbl = doc.add_table(rows=1+len(projects), cols=len(cols_def))
    tbl.style = "Table Grid"; tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for j, (th, en, wcm) in enumerate(cols_def):
        cell = tbl.rows[0].cells[j]
        _cell_bg(cell, C["navy"]); _cell_borders(cell, color=C["white"], size=4)
        _cell_margins(cell, 80,80,90,90); _col_width(cell, int(Cm(wcm).twips))
        hp = _cp(cell)
        hp.alignment = (WD_ALIGN_PARAGRAPH.CENTER if j in (0,6)
                        else WD_ALIGN_PARAGRAPH.LEFT)
        _run_th(hp, th+"\n", size=8.5, bold=True, color_key="white")
        _run_en(hp, en, size=7,
                color_key=None)
        # steelblue for header sub
        hp.runs[-1].font.color.rgb = RGBColor(0x90,0xB8,0xE0)

    aligns = ([WD_ALIGN_PARAGRAPH.CENTER] +
              [WD_ALIGN_PARAGRAPH.LEFT]*5 +
              [WD_ALIGN_PARAGRAPH.RIGHT])
    for i, proj in enumerate(projects):
        bg = C["rowalt"] if i%2==0 else C["white"]
        vals = [
            str(proj.get("no", i+1)),
            (proj.get("name_th",""), proj.get("name_en","")),
            (proj.get("customer_th",""), proj.get("customer_en","")),
            (proj.get("main_contact","—"), ""),
            (proj.get("location","—"), ""),
            (proj.get("desc_th",""), proj.get("desc_en","")),
            ("฿ "+proj.get("value","—"), ""),
        ]
        for j, (val, aln, (_,_,wcm)) in enumerate(zip(vals, aligns, cols_def)):
            cell = tbl.rows[1+i].cells[j]
            _cell_bg(cell, bg); _cell_borders(cell, color="D8E6F2")
            _cell_margins(cell, 70,70,90,90); _col_width(cell, int(Cm(wcm).twips))
            dp = _cp(cell); dp.alignment = aln
            if isinstance(val, tuple):
                th_txt, en_txt = val
                if th_txt:
                    _run_th(dp, th_txt, size=9,
                            bold=(j in (1,)),
                            color_key="navy" if j==1 else "textdark")
                if en_txt:
                    dp.add_run("\n")
                    _run_en(dp, en_txt, size=7.5, color_key="midgray")
            else:
                _run_th(dp, val, size=9,
                        bold=(j in (0,6)),
                        color_key="blue" if j==0 else "navy" if j==6 else "textdark")

    _footer_bar(doc, co, 3)
    doc.add_page_break()


def _build_case_study(doc: Document, co: dict, proj: dict, page_num: int):
    _page_top_rule(doc)
    _running_header(doc, co)

    p = doc.add_paragraph()
    _para_space(p, before=20, after=60)
    _run_th(p, "กรณีศึกษาโครงการ  ", size=15, bold=True, color_key="navy")
    _run_en(p, "CASE STUDY FOR PROJECTS REFERENCE", size=9, color_key="midgray")
    _bottom_rule(p, color=C["blue"], size=12)
    _spacer(doc, 8)

    # Project title banner
    bt = doc.add_table(rows=1, cols=1); bt.alignment = WD_TABLE_ALIGNMENT.LEFT
    bc = bt.rows[0].cells[0]
    _cell_bg(bc, C["ltblue"]); _cell_borders(bc, color=C["navy"], size=12)
    _cell_margins(bc, 110,110,200,140); _col_width(bc, int(Cm(18.0).twips))
    lp = _cp(bc)
    _run_th(lp, f"โครงการที่ {proj.get('no','—')}  ·  {proj.get('name_th','')}",
            size=12, bold=True, color_key="navy")
    lp.add_run("\n")
    _run_en(lp, f"Project No.{proj.get('no','—')}  ·  {proj.get('name_en','')}",
            size=9, italic=True, color_key="midgray")

    _spacer(doc, 12)
    _section_header(doc, "สรุปข้อมูลโครงการ", "Project Summary")
    _spacer(doc, 6)

    rows_data = [
        ("ชื่อโครงการ\nProject Name",
         proj.get("name_th","") + "\n" + proj.get("name_en","")),
        ("ลูกค้า / เจ้าของโครงการ\nCustomer / Owner",
         proj.get("customer_th","") + "\n" + proj.get("customer_en","")),
        ("ผู้รับช่วงงาน\nMain Contact",
         proj.get("main_contact","—")),
        ("รายละเอียดงาน\nJob Description",
         proj.get("desc_th","") + "\n" + proj.get("desc_en","")),
        ("สถานที่ดำเนินงาน\nSite Location",
         proj.get("location","—")),
        ("ระยะเวลาดำเนินงาน\nProject Period",
         proj.get("period","—")),
        ("มูลค่างาน\nContract Value",
         f"฿  {proj.get('value','—')}  บาท"),
    ]
    lw = int(Cm(4.5).twips); vw = int(Cm(13.5).twips)
    st2 = doc.add_table(rows=len(rows_data), cols=2)
    st2.style = "Table Grid"; st2.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (lbl, val) in enumerate(rows_data):
        lc2 = st2.rows[i].cells[0]; vc2 = st2.rows[i].cells[1]
        is_val = lbl.startswith("มูลค่า")
        _cell_bg(lc2, "092F63" if is_val else C["navy"])
        _cell_bg(vc2, "EDF5FF" if is_val else (C["rowalt"] if i%2==0 else C["white"]))
        _cell_borders(lc2, color=C["white"], size=4)
        _cell_borders(vc2, color="D8E6F2")
        _cell_margins(lc2, 80,80,120,80); _cell_margins(vc2, 80,80,120,80)
        _col_width(lc2, lw); _col_width(vc2, vw)
        lp3 = _cp(lc2)
        for k, ln in enumerate(lbl.split("\n")):
            _run_th(lp3, ln+("\n" if k==0 and "\n" in lbl else ""),
                    size=9 if k==0 else 7.5, bold=(k==0),
                    color_key="white" if k==0 else None)
            if k==1: lp3.runs[-1].font.color.rgb = RGBColor(0x90,0xB8,0xE0)
        vp3 = _cp(vc2)
        for k, ln in enumerate(val.split("\n")):
            _run_th(vp3, ln+("\n" if k < len(val.split("\n"))-1 else ""),
                    size=12 if is_val and k==0 else (9.5 if k==0 else 8),
                    bold=is_val and k==0,
                    color_key="navy" if is_val else ("textdark" if k==0 else "midgray"))

    _spacer(doc, 14)
    _section_header(doc, "ภาพถ่ายผลงานโครงการ", "Project Photography")
    _spacer(doc, 8)

    photo_paths = [PHOTOS_DIR/fn for fn in proj.get("photos",[])
                   if (PHOTOS_DIR/fn).exists()]
    if photo_paths:
        _photo_grid(doc, photo_paths, img_width_cm=4.0, cols=4)
    else:
        # Empty placeholder rows — no dummy text, just light boxes
        for _ in range(2):
            tbl = doc.add_table(rows=1, cols=4); tbl.style="Table Grid"
            for j in range(4):
                cell = tbl.rows[0].cells[j]
                _cell_bg(cell, C["ltblue"]); _cell_borders(cell, color="C8D8E8")
                _cell_margins(cell); _col_width(cell, int(Cm(4.3).twips))
                pp = _cp(cell); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _run_en(pp, "(photo)", size=7, italic=True, color_key="midgray")
            _spacer(doc, 6)

    _footer_bar(doc, co, page_num)
    doc.add_page_break()


def _build_custom_section(doc: Document, co: dict, section: dict, page_num: int):
    """
    Render a user-created custom section.
    section = {
        'id', 'title_th', 'title_en',
        'desc_th', 'desc_en',
        'photos': [filename, ...]
    }
    """
    _page_top_rule(doc)
    _running_header(doc, co)

    # Section page title
    p = doc.add_paragraph()
    _para_space(p, before=20, after=60)
    _run_th(p, section.get("title_th","") + "  ", size=15, bold=True, color_key="navy")
    _run_en(p, section.get("title_en","").upper(), size=9, color_key="midgray")
    _bottom_rule(p, color=C["blue"], size=12)
    _spacer(doc, 12)

    # Description block (if provided)
    desc_th = section.get("desc_th","").strip()
    desc_en = section.get("desc_en","").strip()
    if desc_th or desc_en:
        dt = doc.add_table(rows=1, cols=1); dt.alignment = WD_TABLE_ALIGNMENT.LEFT
        dc = dt.rows[0].cells[0]
        _cell_bg(dc, "F4F8FC"); _cell_borders(dc, color=C["blue"], size=6)
        _cell_margins(dc, 120,120,160,160); _col_width(dc, int(Cm(18.0).twips))
        dp = _cp(dc)
        if desc_th:
            _run_th(dp, desc_th+("\n\n" if desc_en else ""),
                    size=10, color_key="textdark")
        if desc_en:
            _run_en(dp, desc_en, size=9, italic=True, color_key="midgray")
        _spacer(doc, 14)

    # Photo gallery
    photo_paths = [PHOTOS_DIR/fn for fn in section.get("photos",[])
                   if (PHOTOS_DIR/fn).exists()]
    if photo_paths:
        _section_header(doc, "ภาพถ่าย", "Photography")
        _spacer(doc, 8)
        _photo_grid(doc, photo_paths, img_width_cm=4.0, cols=4)

    _footer_bar(doc, co, page_num)
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX — MAIN GENERATOR
# ══════════════════════════════════════════════════════════════════════════════

def generate_docx(co: dict, projects: list, sections: list) -> bytes:
    doc = Document()

    # A4 page setup
    for sec in doc.sections:
        sec.page_width    = Cm(21.0)
        sec.page_height   = Cm(29.7)
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(2.0)

    # Default paragraph style — TH Sarabun New + no spacing
    ns = doc.styles["Normal"]
    ns.font.name = THAI_FONT
    ns.font.size = Pt(10)
    # Also set via XML to cover all scripts
    ns_rPr = ns.element.find(qn("w:rPr"))
    ns_rPr = ns_rPr if ns_rPr is not None else OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"):
        rFonts.set(qn(attr), THAI_FONT)
    ns_rPr.insert(0, rFonts)

    _build_cover(doc, co)
    _build_company_info(doc, co)
    _build_project_reference(doc, co, projects)

    page = 4
    for proj in projects:
        _build_case_study(doc, co, proj, page)
        page += 1

    for sect in sections:
        _build_custom_section(doc, co, sect, page)
        page += 1

    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()

    # Save timestamped copy
    ts  = datetime.now().strftime("%Y%m%d_%H%M%S")
    (OUTPUT_DIR / f"DKB_Profile_{ts}.docx").write_bytes(raw)

    return raw


# ══════════════════════════════════════════════════════════════════════════════
#  STREAMLIT  UI
# ══════════════════════════════════════════════════════════════════════════════

CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Sarabun:wght@300;400;500;600;700&display=swap');
html,body,[class*="css"]{font-family:'Sarabun',sans-serif!important;}

[data-testid="stSidebar"]{
  background:linear-gradient(180deg,#062A5A 0%,#0A3C7A 55%,#0D4A8E 100%)!important;
}
[data-testid="stSidebar"] *          {color:#E8F0FA!important;}
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3         {color:#fff!important;font-weight:700!important;}
[data-testid="stSidebar"] hr         {border-color:rgba(255,255,255,.2)!important;}

.main .block-container{max-width:1000px;padding-top:1.5rem;}

.app-header{
  background:linear-gradient(135deg,#0A3C7A 0%,#2979C8 100%);
  border-radius:12px;padding:22px 28px;margin-bottom:20px;
}
.app-header .tag{font-size:10px;letter-spacing:3px;color:#90C8F0;
  text-transform:uppercase;margin-bottom:4px;}
.app-header h1{font-size:22px;font-weight:700;color:#fff;margin:0;line-height:1.3;}
.app-header p{font-size:12px;color:#B3D1F0;margin:4px 0 0;}

.val-badge{display:inline-block;background:#0A3C7A;color:#fff;
  border-radius:4px;padding:2px 10px;font-size:12px;font-weight:600;}
.sec-title{font-size:18px;font-weight:700;color:#0A3C7A;
  border-bottom:2px solid #2979C8;padding-bottom:6px;margin-bottom:14px;}
.info-row{display:flex;gap:6px;align-items:baseline;margin:3px 0;}
.info-label{font-size:11px;color:#8AA0B8;min-width:120px;}
.info-value{font-size:12px;color:#1A2E45;font-weight:500;}

div[data-testid="stForm"]{background:#F9FBFD;border:1px solid #D8E6F2;
  border-radius:10px;padding:20px;}
.stTextInput>label,.stTextArea>label,.stFileUploader>label{
  font-weight:600!important;color:#0A3C7A!important;font-size:13px!important;}
.stTabs [data-baseweb="tab"]{font-size:14px;font-weight:600;}
.stTabs [aria-selected="true"]{color:#0A3C7A!important;}

/* Primary generate button */
.stButton>button[kind="primary"]{
  background:linear-gradient(135deg,#0A3C7A,#2979C8)!important;
  color:#fff!important;font-size:15px!important;font-weight:700!important;
  padding:14px 28px!important;border-radius:8px!important;
  border:none!important;width:100%!important;
}
.stDownloadButton>button{
  background:linear-gradient(135deg,#1A6A50,#2EB382)!important;
  color:#fff!important;font-size:14px!important;font-weight:600!important;
  padding:12px 24px!important;border-radius:8px!important;
  border:none!important;width:100%!important;
}
/* Custom section card */
.cs-card{background:#fff;border:1px solid #D8E6F2;border-radius:10px;
  padding:14px 18px;margin:8px 0;box-shadow:0 2px 8px rgba(10,60,122,.05);}
</style>
"""


# ─── sidebar ─────────────────────────────────────────────────────────────────
def _sidebar(co: dict):
    st.sidebar.markdown("### ⚡ DKB POWER ENGINEERING")
    st.sidebar.markdown(f"**{co['name_th']}**")
    st.sidebar.markdown("---")
    for lbl, val in [
        ("📋 เลขทะเบียน",      co["reg_no"]),
        ("📅 วันที่จดทะเบียน",  co["reg_date_th"]),
        ("💰 ทุนจดทะเบียน",     co["capital_th"]),
        ("📞 โทรศัพท์",         co["tel"]),
        ("✉️ อีเมล",             co["email"]),
    ]:
        st.sidebar.markdown(f"**{lbl}**  {val}")
    st.sidebar.markdown("---")
    st.sidebar.markdown(
        f"<small style='opacity:.5'>DKB Profile Dashboard v3.0 · {co['year']}</small>",
        unsafe_allow_html=True)


# ─── Tab 1: Add project ───────────────────────────────────────────────────────
def _tab_add_project(projects: list):
    st.markdown('<div class="sec-title">➕ เพิ่มโครงการใหม่ / Add New Project</div>',
                unsafe_allow_html=True)
    with st.form("add_proj", clear_on_submit=True):
        c1, c2 = st.columns(2)
        with c1:
            name_th       = st.text_input("ชื่อโครงการ (ไทย) *")
            customer_th   = st.text_input("ลูกค้า / ผู้ว่าจ้าง (ไทย) *")
            main_contact  = st.text_input("ผู้รับช่วงงาน / Main Contact",
                                          placeholder="เช่น บ.ณัฐกิจการช่าง")
            location      = st.text_input("สถานที่ / Location")
            desc_th       = st.text_area("รายละเอียดงาน (ไทย) *", height=80)
        with c2:
            name_en       = st.text_input("Project Name (EN) *")
            customer_en   = st.text_input("Customer / Owner (EN) *")
            period        = st.text_input("ช่วงเวลา / Period")
            value         = st.text_input("มูลค่างาน (THB) *")
            desc_en       = st.text_area("Job Description (EN) *", height=80)

        uploaded = st.file_uploader(
            "📸 อัพโหลดรูปภาพโครงการ (JPG/PNG/WEBP)",
            accept_multiple_files=True, type=["jpg","jpeg","png","webp"],
            label_visibility="visible")
        if uploaded:
            cols = st.columns(min(len(uploaded), 4))
            for i, uf in enumerate(uploaded[:4]):
                with cols[i]:
                    st.image(Image.open(uf), width="stretch")
                    uf.seek(0)

        sub = st.form_submit_button("💾 บันทึกโครงการ", use_container_width=True)

    if sub:
        errs = [f for f,v in [("ชื่อ(ไทย)",name_th),("Name(EN)",name_en),
                               ("ลูกค้า(ไทย)",customer_th),("Customer(EN)",customer_en),
                               ("รายละเอียด(ไทย)",desc_th),("Description(EN)",desc_en),
                               ("มูลค่า",value)] if not v.strip()]
        if errs:
            st.error("⚠️ กรุณากรอก: " + ", ".join(errs)); return None
        photos  = save_photos(uploaded)
        new_proj = {
            "id": f"proj_{uuid.uuid4().hex[:8]}",
            "no": next_no(projects),
            "name_th": name_th.strip(), "name_en": name_en.strip(),
            "customer_th": customer_th.strip(), "customer_en": customer_en.strip(),
            "main_contact": main_contact.strip() or "—",
            "location": location.strip(),
            "desc_th": desc_th.strip(), "desc_en": desc_en.strip(),
            "period": period.strip(), "value": value.strip(),
            "photos": photos, "created_at": datetime.now().isoformat(),
        }
        updated = projects + [new_proj]
        save_projects(updated)
        st.success(f"✅ บันทึกโครงการที่ {new_proj['no']}: {new_proj['name_th']}")
        if photos: st.info(f"📸 บันทึก {len(photos)} รูปภาพ")
        st.balloons()
        return updated
    return None


# ─── Tab 2: View projects ─────────────────────────────────────────────────────
def _tab_view_projects(projects: list):
    st.markdown(
        f'<div class="sec-title">📋 โครงการทั้งหมด '
        f'<span style="font-size:14px;color:#2979C8;">({len(projects)} โครงการ)</span></div>',
        unsafe_allow_html=True)
    if not projects:
        st.info("ยังไม่มีโครงการ"); return None
    for proj in projects:
        with st.expander(f"#{proj.get('no')}  ·  {proj.get('name_th','')}", expanded=False):
            ci, cp = st.columns([3,1])
            with ci:
                for lbl, val in [
                    ("ชื่อ EN",       proj.get("name_en","—")),
                    ("ลูกค้า",        f"{proj.get('customer_th','—')}  ·  {proj.get('customer_en','—')}"),
                    ("ผู้รับช่วงงาน", proj.get("main_contact","—")),
                    ("สถานที่",       proj.get("location","—")),
                    ("รายละเอียด",    f"{proj.get('desc_th','—')} / {proj.get('desc_en','—')}"),
                    ("ช่วงเวลา",      proj.get("period","—")),
                ]:
                    st.markdown(
                        f'<div class="info-row"><span class="info-label">{lbl}</span>'
                        f'<span class="info-value">{val}</span></div>',
                        unsafe_allow_html=True)
                st.markdown(f'มูลค่า: <span class="val-badge">฿ {proj.get("value","—")} บาท</span>',
                            unsafe_allow_html=True)
            with cp:
                photos = proj.get("photos",[])
                if photos and (PHOTOS_DIR/photos[0]).exists():
                    st.image(str(PHOTOS_DIR/photos[0]), width="stretch")
                st.caption(f"📸 {len(photos)} รูป")
            if len(photos)>1:
                tc = st.columns(min(len(photos),4))
                for k,fn in enumerate(photos[:8]):
                    fp = PHOTOS_DIR/fn
                    if fp.exists():
                        with tc[k%4]: st.image(str(fp), width="stretch")

            # Add more photos
            with st.expander("➕ เพิ่มรูปภาพ"):
                more = st.file_uploader("เลือกรูปเพิ่ม", accept_multiple_files=True,
                                        type=["jpg","jpeg","png","webp"],
                                        key=f"more_{proj['id']}")
                if st.button("บันทึกรูปเพิ่ม", key=f"sv_{proj['id']}"):
                    if more:
                        fns = save_photos(more)
                        for p2 in projects:
                            if p2["id"]==proj["id"]:
                                p2["photos"] = p2.get("photos",[])+fns
                        save_projects(projects)
                        st.success(f"เพิ่ม {len(fns)} รูปแล้ว"); st.rerun()

            if st.button("🗑️ ลบโครงการนี้", key=f"del_{proj['id']}", type="secondary"):
                rem = [p for p in projects if p["id"]!=proj["id"]]
                for k,p in enumerate(rem): p["no"]=k+1
                save_projects(rem); st.session_state.projects=rem
                st.warning(f"ลบ '{proj['name_th']}' แล้ว"); st.rerun()
    return None


# ─── Tab 3: Custom sections ───────────────────────────────────────────────────
def _tab_custom_sections(sections: list):
    st.markdown('<div class="sec-title">📂 Custom Sections (Preventive Maintenance, Construction, etc.)</div>',
                unsafe_allow_html=True)
    st.info(
        "เพิ่มส่วนพิเศษลงใน Company Profile เช่น งาน Preventive Maintenance, "
        "โครงการระหว่างก่อสร้าง, หรือ Factory References — "
        "แต่ละส่วนมีชื่อ คำอธิบาย และรูปภาพของตัวเอง",
        icon="💡")

    # Show existing sections
    for i, sec in enumerate(sections):
        with st.expander(
            f"📌 Section {i+1}:  {sec.get('title_th','')}  /  {sec.get('title_en','')}",
            expanded=False):
            st.markdown(f"**คำอธิบาย (ไทย):** {sec.get('desc_th','—')}")
            st.markdown(f"**Description (EN):** {sec.get('desc_en','—')}")
            photos = sec.get("photos",[])
            st.caption(f"📸 {len(photos)} รูปภาพ")
            if photos:
                tc = st.columns(min(len(photos),4))
                for k,fn in enumerate(photos[:8]):
                    fp = PHOTOS_DIR/fn
                    if fp.exists():
                        with tc[k%4]: st.image(str(fp), width="stretch")

            # Add photos to section
            with st.expander("➕ เพิ่มรูปภาพ"):
                more = st.file_uploader("เลือกรูป", accept_multiple_files=True,
                                        type=["jpg","jpeg","png","webp"],
                                        key=f"sph_{sec['id']}")
                if st.button("บันทึกรูปเพิ่ม", key=f"svsph_{sec['id']}"):
                    if more:
                        fns = save_photos(more)
                        sec["photos"] = sec.get("photos",[])+fns
                        save_sections(sections); st.success(f"เพิ่ม {len(fns)} รูป"); st.rerun()

            if st.button("🗑️ ลบ Section นี้", key=f"dsec_{sec['id']}", type="secondary"):
                sections.remove(sec); save_sections(sections)
                st.session_state.sections = sections
                st.warning("ลบ Section แล้ว"); st.rerun()

    st.markdown("---")
    st.markdown("#### ➕ เพิ่ม Section ใหม่ / Add New Section")

    # Preset templates for quick start
    presets = {
        "(เลือก Template)": ("","","",""),
        "Preventive Maintenance": (
            "งาน Preventive Maintenance", "Preventive Maintenance",
            "งานบำรุงรักษาระบบไฟฟ้าเชิงป้องกัน",
            "Scheduled preventive maintenance of electrical systems"),
        "Project under Construction": (
            "โครงการระหว่างก่อสร้าง", "Project under Construction",
            "โครงการที่อยู่ระหว่างดำเนินการก่อสร้าง",
            "Projects currently under construction"),
        "Factory References": (
            "อ้างอิงโครงการโรงงาน", "Factory References",
            "ผลงานการติดตั้งระบบในโรงงานอุตสาหกรรม",
            "Industrial factory electrical installation references"),
        "Solar / Renewable Energy": (
            "โครงการพลังงานแสงอาทิตย์", "Solar & Renewable Energy Projects",
            "งานติดตั้งระบบโซลาร์เซลล์และพลังงานทดแทน",
            "Solar cell and renewable energy installation projects"),
        "Condominium / Residential": (
            "งานอาคารพักอาศัย", "Condominium & Residential Projects",
            "งานติดตั้งระบบไฟฟ้าในอาคารชุดและที่พักอาศัย",
            "Electrical system installation in condominium and residential buildings"),
    }
    preset = st.selectbox("🗂️ ใช้ Template สำเร็จรูป", list(presets.keys()))
    pt = presets[preset]

    with st.form("add_section", clear_on_submit=True):
        c1, c2 = st.columns(2)
        with c1:
            title_th = st.text_input("ชื่อส่วน (ไทย) *",
                                     value=pt[0] if preset!="(เลือก Template)" else "")
            desc_th  = st.text_area("คำอธิบาย (ไทย)", height=80,
                                    value=pt[2] if preset!="(เลือก Template)" else "")
        with c2:
            title_en = st.text_input("Section Title (EN) *",
                                     value=pt[1] if preset!="(เลือก Template)" else "")
            desc_en  = st.text_area("Description (EN)", height=80,
                                    value=pt[3] if preset!="(เลือก Template)" else "")

        uploaded = st.file_uploader(
            "📸 อัพโหลดรูปภาพ (สามารถเพิ่มเติมภายหลังได้)",
            accept_multiple_files=True, type=["jpg","jpeg","png","webp"],
            label_visibility="visible")

        sub = st.form_submit_button("💾 บันทึก Section", use_container_width=True)

    if sub:
        if not title_th.strip() or not title_en.strip():
            st.error("⚠️ กรุณากรอกชื่อ Section"); return None
        photos = save_photos(uploaded)
        new_sec = {
            "id"       : f"sec_{uuid.uuid4().hex[:8]}",
            "title_th" : title_th.strip(),
            "title_en" : title_en.strip(),
            "desc_th"  : desc_th.strip(),
            "desc_en"  : desc_en.strip(),
            "photos"   : photos,
            "created_at": datetime.now().isoformat(),
        }
        updated = sections + [new_sec]
        save_sections(updated)
        st.success(f"✅ เพิ่ม Section: {new_sec['title_th']}")
        return updated
    return None


# ─── Tab 4: Company info ──────────────────────────────────────────────────────
def _tab_company(co: dict):
    st.markdown('<div class="sec-title">🏢 ข้อมูลบริษัท / Company Information</div>',
                unsafe_allow_html=True)
    with st.form("edit_co"):
        c1, c2 = st.columns(2)
        with c1:
            name_th    = st.text_input("ชื่อบริษัท (ไทย)",    co["name_th"])
            reg_no     = st.text_input("เลขทะเบียนนิติบุคคล", co["reg_no"])
            reg_th     = st.text_input("วันที่จดทะเบียน (ไทย)",co["reg_date_th"])
            capital_th = st.text_input("ทุนจดทะเบียน (ไทย)",   co["capital_th"])
            tel        = st.text_input("โทรศัพท์",             co["tel"])
        with c2:
            name_en    = st.text_input("Company Name (EN)",    co["name_en"])
            year       = st.text_input("ปี / Year",            co["year"])
            reg_en     = st.text_input("Date of Incorporation (EN)", co["reg_date_en"])
            capital_en = st.text_input("Registered Capital (EN)",    co["capital_en"])
            email      = st.text_input("อีเมล",                co["email"])
        addr_th    = st.text_area("ที่อยู่ (ไทย)",  co["address_th"], height=60)
        addr_en    = st.text_area("Address (EN)",   co["address_en"], height=60)
        hist_th    = st.text_area("ประวัติบริษัท (ไทย)", co["history_th"], height=120)
        hist_en    = st.text_area("Company History (EN)",   co["history_en"], height=120)

        logo_file  = st.file_uploader("โลโก้บริษัท (PNG/JPG)", type=["png","jpg","jpeg"])
        if co.get("logo_file") and (ASSETS_DIR/co["logo_file"]).exists():
            st.image(str(ASSETS_DIR/co["logo_file"]), width=180)

        saved = st.form_submit_button("💾 บันทึกข้อมูลบริษัท", use_container_width=True)

    if saved:
        logo_fn = co.get("logo_file","")
        if logo_file: logo_fn = save_logo(logo_file)
        updated = {**co,
            "name_th":name_th,"name_en":name_en,"reg_no":reg_no,
            "reg_date_th":reg_th,"reg_date_en":reg_en,
            "capital_th":capital_th,"capital_en":capital_en,
            "address_th":addr_th,"address_en":addr_en,
            "tel":tel,"email":email,"year":year,
            "history_th":hist_th,"history_en":hist_en,
            "logo_file":logo_fn,
        }
        save_company(updated)
        st.success("✅ บันทึกข้อมูลบริษัทสำเร็จ")
        return updated
    return None


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    st.set_page_config(
        page_title="DKB Power Engineering — Profile Dashboard",
        page_icon="⚡", layout="wide",
        initial_sidebar_state="expanded")
    st.markdown(CSS, unsafe_allow_html=True)

    # State init
    if "company"    not in st.session_state: st.session_state.company    = load_company()
    if "projects"   not in st.session_state: st.session_state.projects   = load_projects()
    if "sections"   not in st.session_state: st.session_state.sections   = load_sections()
    if "docx_bytes" not in st.session_state: st.session_state.docx_bytes = None

    co       = st.session_state.company
    projects = st.session_state.projects
    sections = st.session_state.sections

    _sidebar(co)

    # Header
    st.markdown(f"""
    <div class="app-header">
      <div class="tag">Company Profile Dashboard  ·  v3.0</div>
      <h1>⚡ DKB POWER ENGINEERING CO., LTD.</h1>
      <p>บริษัท ดีเคบี เพาเวอร์ เอนจิเนียริ่ง จำกัด  ·
         {len(projects)} โครงการ  ·  {len(sections)} Custom Section(s)</p>
    </div>""", unsafe_allow_html=True)

    tab1, tab2, tab3, tab4 = st.tabs([
        "➕ เพิ่มโครงการ",
        "📋 ดูโครงการทั้งหมด",
        "📂 Custom Sections",
        "🏢 ข้อมูลบริษัท",
    ])

    with tab1:
        r = _tab_add_project(projects)
        if r is not None:
            st.session_state.projects = r; projects = r

    with tab2:
        _tab_view_projects(projects)

    with tab3:
        r = _tab_custom_sections(sections)
        if r is not None:
            st.session_state.sections = r; sections = r

    with tab4:
        r = _tab_company(co)
        if r is not None:
            st.session_state.company = r; co = r; st.rerun()

    # ── Generate bar ─────────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown(
        f'<div style="text-align:center;font-size:16px;font-weight:700;'
        f'color:#0A3C7A;margin-bottom:10px;">'
        f"📄 พร้อมสร้างเอกสาร — {len(projects)} โครงการ + {len(sections)} Section(s)</div>",
        unsafe_allow_html=True)

    gc, dc = st.columns(2)
    with gc:
        if st.button("⚙️ สร้าง Company Profile (.docx)",
                     use_container_width=True, type="primary"):
            with st.spinner("กำลังสร้างเอกสาร…"):
                try:
                    raw = generate_docx(co, projects, sections)
                    st.session_state.docx_bytes = raw
                    st.success("✅ สร้างเอกสารสำเร็จ!")
                except Exception as exc:
                    st.error(f"❌ ข้อผิดพลาด: {exc}"); raise
    with dc:
        if st.session_state.docx_bytes:
            st.download_button(
                label="⬇️ Download Company Profile (.docx)",
                data=st.session_state.docx_bytes,
                file_name=f"DKB_Power_Engineering_Company_Profile_{co['year']}.docx",
                mime="application/vnd.openxmlformats-officedocument"
                     ".wordprocessingml.document",
                use_container_width=True)

    st.markdown(
        f"<div style='text-align:center;margin-top:28px;font-size:11px;color:#AAB8C8;'>"
        f"DKB Power Engineering Co., Ltd.  ·  Profile Dashboard v3.0  ·  {co['year']}</div>",
        unsafe_allow_html=True)


if __name__ == "__main__":
    main()
